from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

import pandas as pd
from sqlalchemy.orm import Session

from backend import models

try:
    from docx import Document
except Exception:
    Document = None


@dataclass
class CurriculumImportStats:
    file_path: str
    program_code: str
    total_rows: int
    parsed_courses: int
    inserted_courses: int
    updated_courses: int
    deleted_courses: int
    protected_courses: int
    inserted_rules: int
    updated_rules: int
    assumptions: list[str]


def _to_float(value) -> float | None:
    if value is None:
        return None
    text = str(value).strip().replace(",", ".")
    if not text:
        return None
    try:
        out = float(text)
    except Exception:
        return None
    return out if out > 0 else None


def _norm(s: str) -> str:
    return " ".join(str(s or "").lower().split())


def _find_col(columns: list[str], keys: list[str]) -> str | None:
    for col in columns:
        n = _norm(col)
        if any(k in n for k in keys):
            return col
    return None


_SPEC_DOCX_TO_DISPLAY: dict[str, str] = {
    "cntt dia hoc": "Công nghệ thông tin Địa học",
    "cong nghe thong tin dia hoc": "Công nghệ thông tin Địa học",
    "khoa hoc may tinh": "Khoa học máy tính",
    "tin hoc kinh te": "Tin học kinh tế",
    "mang may tinh": "Mạng máy tính",
    "he thong thong tin": "Hệ thống thông tin",
    "cong nghe phan mem": "Công nghệ phần mềm",
}

_GROUP_MIN_CREDITS: dict[str, float] = {"A": 6.0, "B": 9.0, "C": 9.0}


def _norm_spec(s: str) -> str:
    import unicodedata as _u
    s = s.replace("đ", "d").replace("Đ", "D")
    return "".join(c for c in _u.normalize("NFKD", s.lower()) if not _u.combining(c)).strip()


# Space-insensitive lookup (handles OCR artifacts like "k inh tế" → "kinh tế")
_SPEC_DOCX_NOSPACE: dict[str, str] = {
    k.replace(" ", ""): v for k, v in _SPEC_DOCX_TO_DISPLAY.items()
}


def _lookup_spec(raw: str) -> str | None:
    n = _norm_spec(raw)
    if n in _SPEC_DOCX_TO_DISPLAY:
        return _SPEC_DOCX_TO_DISPLAY[n]
    return _SPEC_DOCX_NOSPACE.get(n.replace(" ", ""))


def _detect_specialization(text: str) -> str | None:
    """Extract normalised display-name specialization from a heading paragraph."""
    m = re.search(r"Chuyên ngành\s+(.+)", text[:120], re.IGNORECASE)
    if not m:
        return None
    raw = m.group(1).strip()
    # Strip any trailing "Tự chọn …" that got merged into the same paragraph
    raw = re.sub(r"\s+[Tt]ự\s+[Cc]họn.*", "", raw).strip()
    # Text is sometimes repeated 3× separated by two spaces
    raw = raw.split("  ")[0].strip()
    return _lookup_spec(raw)


def _detect_elective_group(text: str) -> str | None:
    """Detect elective group label from a heading/paragraph text.
    Returns e.g. 'A', 'B', 'I', 'II' or None if not an elective group heading.
    """
    t = _norm(text)
    # Normalize Vietnamese diacritics for matching
    import unicodedata
    t = "".join(
        c for c in unicodedata.normalize("NFKD", t.replace("đ", "d"))
        if not unicodedata.combining(c)
    )
    for pattern in [
        r"nhom\s+([a-z]{1,3})\b",
        r"group\s+([a-z]{1,3})\b",
        r"tu\s+chon\s+nhom\s+([a-z]{1,3})\b",
        r"tu\s+chon\s+([a-z]{1,3})\b",
    ]:
        m = re.search(pattern, t)
        if m:
            return m.group(1).upper()
    return None


def _parse_docx(
    file_path: Path,
) -> tuple[list[tuple[str, str, float | None]], list[tuple[str, str | None, str]], int]:
    """Parse a .docx curriculum file.

    Returns:
        courses: list of (code, name, credits)
        elective_mappings: list of (course_code, specialization_display_name | None, group_letter)
            specialization is None for Group A (global/Chung)
        total_rows: int
    """
    if Document is None:
        raise ValueError("python-docx is required to import .docx curriculum files")

    doc = Document(file_path)
    seen: dict[str, tuple[str, str, float | None, str | None]] = {}
    elective_mappings: list[tuple[str, str | None, str]] = []
    rows = 0
    current_group: str | None = None
    current_spec: str | None = None          # for elective section
    current_required_spec: str | None = None  # for required-course section
    in_elective_section: bool = False
    done_processing: bool = False            # True after all curriculum tables are done
    table_index = 0

    for elem in doc.element.body:
        local = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

        if local == "p":
            raw = " ".join(elem.itertext()).strip()
            if raw:
                # Detect start of elective section
                if "danh mục học phần tự chọn" in raw.lower() or "danh muc hoc phan tu chon" in raw.lower():
                    in_elective_section = True
                    current_required_spec = None  # stop tagging required courses
                # Detect end of elective section (new numbered chapter, e.g. "7.4.")
                elif in_elective_section and re.match(r"\d+\.\d+\.", raw):
                    in_elective_section = False
                    done_processing = True    # stop adding courses from subsequent tables
                    current_group = None
                    current_spec = None

                if in_elective_section:
                    spec = _detect_specialization(raw)
                    if spec:
                        current_spec = spec
                    detected = _detect_elective_group(raw)
                    if detected:
                        current_group = detected

        elif local == "tbl":
            table_index += 1
            if done_processing:
                continue  # skip matrix/schedule tables after curriculum listing
            if table_index - 1 >= len(doc.tables):
                continue
            table = doc.tables[table_index - 1]

            for row in table.rows:
                rows += 1
                cells = [" ".join(c.text.replace("\n", " ").split()).strip() for c in row.cells]

                # Detect "A/B/C/D/E/F. Chuyên ngành X" header rows in required-course tables
                if not in_elective_section:
                    spec_detected = False
                    for cell in cells:
                        m = re.match(r"^[A-F]\.\s*(Chuyên ngành\s+.+)", cell, re.IGNORECASE)
                        if m:
                            # Normalise: collapse extra internal whitespace then strip prefix
                            raw_spec = re.sub(r"\s+", " ", m.group(1)).strip()
                            raw_spec = re.sub(r"(?i)^chuyên\s+ngành\s+", "", raw_spec).strip()
                            raw_spec = raw_spec.split("  ")[0].strip()
                            new_spec = _lookup_spec(raw_spec)
                            if new_spec:
                                current_required_spec = new_spec
                            spec_detected = True
                            break
                    if spec_detected:
                        continue  # header row, skip course code search

                # Find 7-digit course code
                code = None
                code_idx = -1
                for i, v in enumerate(cells):
                    if re.fullmatch(r"\d{7}", v):
                        code = v
                        code_idx = i
                        break
                if not code:
                    continue

                # Find course name (first non-numeric cell after code)
                name = None
                for i in range(code_idx + 1, min(len(cells), code_idx + 5)):
                    v = cells[i]
                    if v and not re.fullmatch(r"\d+(?:[\.,]\d+)?", v):
                        name = v
                        break
                if not name:
                    continue

                # Find credits (first numeric value <= 12 after code)
                credits = None
                for i in range(code_idx + 1, min(len(cells), code_idx + 8)):
                    credits = _to_float(cells[i])
                    if credits is not None and credits <= 12:
                        break
                    credits = None

                req_spec = current_required_spec if not in_elective_section else None
                # Don't overwrite an existing entry's required_spec with None
                if code not in seen:
                    seen[code] = (code, name, credits, req_spec)
                else:
                    # Update name/credits if not set; don't downgrade required_spec to None
                    existing = seen[code]
                    seen[code] = (code, name or existing[1], credits if credits is not None else existing[2], existing[3] or req_spec)

                if current_group:
                    elective_mappings.append((code, current_spec, current_group))

    return list(seen.values()), elective_mappings, rows


def _parse_tabular(
    file_path: Path,
) -> tuple[list[tuple[str, str, float | None]], list[tuple[str, str | None, str]], int]:
    """Parse a .xlsx/.csv curriculum file.

    Returns:
        courses: list of (code, name, credits)
        elective_mappings: list of (course_code, specialization | None, group_type)
        total_rows: int
    """
    suffix = file_path.suffix.lower()
    if suffix in {".xlsx", ".xls", ".xlsm"}:
        df = pd.read_excel(file_path, sheet_name=0, engine="openpyxl")
    elif suffix == ".csv":
        df = pd.read_csv(file_path, encoding="utf-8-sig")
    else:
        raise ValueError("Unsupported curriculum file format. Use .xlsx/.xls/.xlsm/.csv/.docx")

    code_col = _find_col(list(df.columns), ["ma mh", "ma hoc phan", "ma mon", "course_code", "code"])
    name_col = _find_col(list(df.columns), ["ten mon", "ten hoc phan", "ten mon hoc", "course_name", "name"])
    credit_col = _find_col(list(df.columns), ["so tin chi", "tin chi", "credits", "credit"])
    group_col = _find_col(list(df.columns), ["nhom", "group", "to", "elective_group"])

    if not code_col or not name_col:
        raise ValueError("Cannot detect required columns: course_code and course_name")

    seen: dict[str, tuple[str, str, float | None]] = {}
    elective_mappings: list[tuple[str, str | None, str]] = []

    for _, row in df.iterrows():
        code = str(row.get(code_col) or "").strip()
        if code.endswith(".0") and code[:-2].isdigit():
            code = code[:-2]
        if not re.fullmatch(r"\d{7}", code):
            continue
        name = str(row.get(name_col) or "").strip()
        if not name:
            continue
        credits = _to_float(row.get(credit_col)) if credit_col else None
        seen[code] = (code, name, credits, None)

        if group_col:
            group_val = str(row.get(group_col) or "").strip()
            if group_val:
                elective_mappings.append((code, None, group_val.upper()))

    return list(seen.values()), elective_mappings, int(len(df.index))


def import_curriculum_from_excel(
    db: Session,
    file_path: str = "data/ChuongTrinhDaoTao.xlsx",
    *,
    replace_existing: bool = True,
    default_program_code: str = "7480201",
    default_specialization: str = "Chung",
) -> CurriculumImportStats:
    resolved = Path(file_path)
    if not resolved.is_absolute():
        resolved = Path.cwd() / resolved
    if not resolved.exists():
        raise FileNotFoundError(f"Curriculum file not found: {resolved}")

    if resolved.suffix.lower() == ".docx":
        parsed, elective_mappings, total_rows = _parse_docx(resolved)
    else:
        parsed, elective_mappings, total_rows = _parse_tabular(resolved)

    # --- Upsert courses ---
    existing_courses = {c.course_code: c for c in db.query(models.Course).all()}
    imported_codes: set[str] = set()
    inserted_courses = 0
    updated_courses = 0
    deleted_courses = 0
    protected_courses = 0

    for entry in parsed:
        code, name, credits = entry[0], entry[1], entry[2]
        req_spec = entry[3] if len(entry) > 3 else None
        imported_codes.add(code)
        existing = existing_courses.get(code)
        if existing:
            existing.course_name = name
            if credits is not None:
                existing.credits = credits
            if req_spec is not None:
                existing.required_specialization = req_spec
            updated_courses += 1
        else:
            db.add(models.Course(course_code=code, course_name=name, credits=credits, required_specialization=req_spec))
            inserted_courses += 1

    if replace_existing:
        referenced_codes = {row[0] for row in db.query(models.UserGrade.course_code).distinct().all()}
        for code, existing in existing_courses.items():
            if code in imported_codes:
                continue
            if code in referenced_codes:
                protected_courses += 1
                continue
            db.delete(existing)
            deleted_courses += 1

    # --- Upsert elective group mappings ---
    inserted_rules = 0
    updated_rules = 0

    if elective_mappings:
        # Flush courses to DB so FK constraints are satisfied
        db.flush()

        if replace_existing:
            # Clear existing mappings and rules for a clean re-import
            db.query(models.CourseElectiveGroup).filter(
                models.CourseElectiveGroup.program_code == default_program_code
            ).delete()
            db.query(models.ElectiveRule).filter(
                models.ElectiveRule.program_code == default_program_code
            ).delete()
            db.flush()

        existing_mappings: set[tuple] = set()
        specs_and_groups_seen: set[tuple[str, str]] = set()

        for course_code, raw_spec, group_type in elective_mappings:
            effective_spec = raw_spec if raw_spec else "Chung"
            key = (course_code, default_program_code, effective_spec, group_type)
            if key not in existing_mappings:
                db.add(models.CourseElectiveGroup(
                    course_code=course_code,
                    program_code=default_program_code,
                    specialization=effective_spec,
                    group_type=group_type,
                ))
                inserted_rules += 1
                existing_mappings.add(key)
            specs_and_groups_seen.add((effective_spec, group_type))

        # Upsert elective_rules based on discovered groups
        for (spec, group) in specs_and_groups_seen:
            min_credits = _GROUP_MIN_CREDITS.get(group, 0.0)
            if min_credits > 0:
                db.add(models.ElectiveRule(
                    program_code=default_program_code,
                    specialization=spec,
                    group_type=group,
                    min_credits_required=min_credits,
                ))

    assumptions: list[str] = []
    if not elective_mappings:
        assumptions.append("Không tìm thấy thông tin nhóm tự chọn trong file. Tiến độ nhóm tự chọn sẽ không hiển thị.")

    return CurriculumImportStats(
        file_path=str(resolved),
        program_code=default_program_code,
        total_rows=total_rows,
        parsed_courses=len(parsed),
        inserted_courses=inserted_courses,
        updated_courses=updated_courses,
        deleted_courses=deleted_courses,
        protected_courses=protected_courses,
        inserted_rules=inserted_rules,
        updated_rules=updated_rules,
        assumptions=assumptions,
    )
